"""One-time setup: writes template.docx next to this file.

The template uses docxtpl's Jinja2 syntax — a {% for %} loop expands the
transfer rows when app.py renders one document per driver. If logo.png
exists next to this file, it is inserted at the top of the template.
Run once (or again whenever logo.png changes):

    python create_template.py
"""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm

HERE = Path(__file__).parent.resolve()
OUT = HERE / "template.docx"
LOGO = HERE / "logo.png"

doc = Document()

if LOGO.exists():
    logo_p = doc.add_paragraph()
    logo_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    logo_p.add_run().add_picture(str(LOGO), width=Cm(5))
    print(f"Inserted logo from {LOGO}")
else:
    print(f"No logo.png found at {LOGO} — skipping (template still valid).")

doc.add_paragraph("{% for driver in drivers %}")
doc.add_paragraph("Şöför: {{driver.name}}")
doc.add_paragraph("")
doc.add_paragraph("{% for row in driver.rows %}")
doc.add_paragraph(
    "{{row.Tarih}} - {{row.Saat}} – {{row.Müşteri}} – {{row.To}} "
    "=> {{row.From}} – {{row.NumberOfPeople}} – {{row.Kişi}}"
)
doc.add_paragraph("{% endfor %}")
doc.add_paragraph("")
doc.add_paragraph("{% endfor %}")

doc.save(str(OUT))
print(f"Wrote {OUT}")
